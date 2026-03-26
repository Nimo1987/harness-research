#!/usr/bin/env python3
"""
DOCX 渲染脚本 — 将 HTML 报告解析并转换为 Word 文档

用法：
  python render_docx.py \
    --input clean_report.html \
    --output output_dir/
"""

import argparse
import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ===== 颜色常量 =====
COLOR_PRIMARY = RGBColor(0x1A, 0x36, 0x5D)
COLOR_ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
COLOR_TEXT = RGBColor(0x2D, 0x37, 0x48)
COLOR_LIGHT = RGBColor(0x71, 0x80, 0x96)
COLOR_CONF_HIGH = RGBColor(0x27, 0x67, 0x49)
COLOR_CONF_MED = RGBColor(0xC0, 0x56, 0x21)
COLOR_CONF_LOW = RGBColor(0xC5, 0x30, 0x30)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def safe_filename(text: str, max_len: int = 30) -> str:
    safe = re.sub(r'[\\/:*?"<>|]', '', text)
    return safe[:max_len].strip()


def extract_title(html: str) -> str:
    match = re.search(r'<h1[^>]*>([^<]+)</h1>', html)
    return match.group(1).strip() if match else "深度调研报告"


def setup_styles(doc: Document):
    """设置文档默认样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = COLOR_TEXT
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def process_element(doc: Document, element):
    """递归处理 HTML 元素"""
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            doc.add_paragraph(text)
        return

    tag = element.name
    if tag is None:
        return

    if tag == 'h1':
        p = doc.add_heading(element.get_text(strip=True), level=0)
        style_heading(p, 24, COLOR_PRIMARY)

    elif tag == 'h2':
        p = doc.add_heading(element.get_text(strip=True), level=1)
        style_heading(p, 16, COLOR_PRIMARY)

    elif tag == 'h3':
        p = doc.add_heading(element.get_text(strip=True), level=2)
        style_heading(p, 13, COLOR_ACCENT)

    elif tag == 'p':
        p = doc.add_paragraph()
        process_inline(p, element)

    elif tag == 'table':
        process_table(doc, element)

    elif tag == 'blockquote':
        p = doc.add_paragraph(element.get_text(strip=True))
        p.paragraph_format.left_indent = Cm(1.5)
        for run in p.runs:
            run.italic = True

    elif tag in ('ol', 'ul'):
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(li.get_text(strip=True), style='List Bullet')

    elif tag == 'figure':
        cap = element.find('figcaption')
        if cap:
            p = doc.add_paragraph(cap.get_text(strip=True))
            if p.runs:
                p.runs[0].bold = True

    elif tag == 'div':
        classes = element.get('class', [])
        if isinstance(classes, str):
            classes = classes.split()
        if 'source-note' in classes:
            p = doc.add_paragraph(element.get_text(strip=True))
            for run in p.runs:
                run.font.size = Pt(8)
                run.italic = True
                run.font.color.rgb = COLOR_LIGHT
        elif 'report-meta' in classes:
            p = doc.add_paragraph(element.get_text(strip=True))
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_LIGHT
        else:
            for child in element.children:
                process_element(doc, child)
    else:
        for child in element.children:
            process_element(doc, child)


def process_inline(paragraph, element):
    """处理段落内行内元素"""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                paragraph.add_run(text)
        elif child.name in ('strong', 'b'):
            run = paragraph.add_run(child.get_text())
            run.bold = True
            run.font.color.rgb = COLOR_PRIMARY
        elif child.name in ('em', 'i'):
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'sup':
            run = paragraph.add_run(child.get_text())
            run.font.size = Pt(7)
            run.font.color.rgb = COLOR_ACCENT
            run.font.superscript = True
        elif child.name == 'span':
            classes = child.get('class', [])
            if isinstance(classes, str):
                classes = classes.split()
            run = paragraph.add_run(child.get_text())
            if 'confidence' in classes:
                run.font.size = Pt(8)
                if 'high' in classes:
                    run.font.color.rgb = COLOR_CONF_HIGH
                elif 'medium' in classes:
                    run.font.color.rgb = COLOR_CONF_MED
                elif 'low' in classes:
                    run.font.color.rgb = COLOR_CONF_LOW
        elif child.name == 'a':
            run = paragraph.add_run(child.get_text())
            run.font.color.rgb = COLOR_ACCENT
        else:
            paragraph.add_run(child.get_text())


def process_table(doc: Document, table_el):
    """将 HTML 表格转为 Word 表格"""
    caption = table_el.find('caption')
    if caption:
        p = doc.add_paragraph(caption.get_text(strip=True))
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = COLOR_PRIMARY

    rows_data = []
    thead = table_el.find('thead')
    tbody = table_el.find('tbody')

    if thead:
        for tr in thead.find_all('tr'):
            cells = [c.get_text(strip=True) for c in tr.find_all(['th', 'td'])]
            rows_data.append(('header', cells))

    body = tbody if tbody else table_el
    for tr in body.find_all('tr'):
        if thead and tr.parent == thead:
            continue
        cells = [c.get_text(strip=True) for c in tr.find_all(['td', 'th'])]
        if cells:
            rows_data.append(('body', cells))

    if not rows_data:
        return

    max_cols = max(len(r[1]) for r in rows_data)
    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (row_type, cells) in enumerate(rows_data):
        row = table.rows[i]
        for j, text in enumerate(cells):
            if j < max_cols:
                cell = row.cells[j]
                cell.text = text
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.space_before = Pt(2)
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if row_type == 'header':
                            run.bold = True
                            run.font.color.rgb = COLOR_WHITE

        if row_type == 'header':
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                shading = tc_pr.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): '1a365d',
                })
                tc_pr.append(shading)

    doc.add_paragraph()


def style_heading(paragraph, size: int, color: RGBColor):
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        html = f.read()

    title = extract_title(html)

    doc = Document()
    setup_styles(doc)

    soup = BeautifulSoup(html, "html.parser")
    for element in soup.children:
        process_element(doc, element)

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"深度调研报告_{safe_filename(title)}.docx"
    doc.save(str(docx_path))

    print(f"[RENDER-DOCX] 完成 → {docx_path}")


if __name__ == "__main__":
    main()
